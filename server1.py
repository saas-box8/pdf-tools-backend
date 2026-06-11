"""
Word (DOCX/DOC/ODT/RTF) → PDF  —  LibreOffice headless backend
Runs on Render free plan (Python 3.11).

Strategy
---------
1. PRIMARY: LibreOffice headless  — pixel-perfect output, handles all .docx
   formatting: fonts, images, tables, headers/footers, track changes, etc.
2. FALLBACK: ReportLab            — pure Python, no binary needed; used only
   when LibreOffice is not available (shouldn't happen on Render).

Render free-plan constraints handled:
- Ephemeral disk: all temp files go to /tmp and are cleaned up on every request.
- Single-process RAM limit: LibreOffice gets a unique --UserInstallation per
  request so concurrent requests never collide on lock files.
- 512 MB RAM: LibreOffice Writer uses ~150 MB; well within limits.
- No persistent storage: nothing is written outside /tmp.
- Timeout: gunicorn --timeout 120 s; LibreOffice conversion of a 100-page DOCX
  typically finishes in under 30 s.
"""

from __future__ import annotations

import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
import time
from pathlib import Path
from typing import Optional
from xml.sax.saxutils import escape

from flask import Flask, Response, jsonify, request, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename

# python-docx  (fallback path only)
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# reportlab  (fallback path only)
from reportlab.lib import colors as rl_colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# ── Logging ───────────────────────────────────────────────────────────────────

logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO"),
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger(__name__)

# ── App setup ─────────────────────────────────────────────────────────────────

app = Flask(__name__)
CORS(app)
app.config["MAX_CONTENT_LENGTH"] = int(os.getenv("MAX_UPLOAD_MB", "100")) * 1024 * 1024

# ── Accepted MIME / extension map ─────────────────────────────────────────────

ACCEPTED_EXTENSIONS = {".docx", ".doc", ".odt", ".rtf"}
ACCEPTED_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/vnd.oasis.opendocument.text",
    "application/rtf",
    "text/rtf",
}

# ═══════════════════════════════════════════════════════════════════════════════
# PRIMARY ENGINE — LibreOffice headless
# ═══════════════════════════════════════════════════════════════════════════════

_SOFFICE_BIN: Optional[str] = shutil.which("libreoffice") or shutil.which("soffice")


def _libreoffice_available() -> bool:
    return _SOFFICE_BIN is not None


def _convert_with_libreoffice(
    input_path: str,
    timeout: int = 90,
) -> bytes:
    """
    Convert a Word/ODT/RTF file to PDF using LibreOffice headless.

    Each call gets an isolated UserInstallation directory so concurrent
    requests can never fight over the same lock files.

    Returns the PDF bytes on success.
    Raises RuntimeError on failure.
    """
    out_dir = tempfile.mkdtemp(prefix="lo_out_")
    profile = tempfile.mkdtemp(prefix="lo_prof_")

    try:
        env = os.environ.copy()
        # LibreOffice needs HOME when running as root / in containers
        env["HOME"] = profile

        cmd = [
            _SOFFICE_BIN,
            "--headless",
            "--norestore",
            "--nofirststartwizard",
            "--nologo",
            # Isolated profile — critical for concurrency
            f"-env:UserInstallation=file://{profile}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            out_dir,
            input_path,
        ]

        logger.info("LibreOffice: converting %s", Path(input_path).name)
        t0 = time.monotonic()

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=timeout,
            env=env,
        )

        elapsed = time.monotonic() - t0
        logger.info("LibreOffice: finished in %.2f s (rc=%d)", elapsed, result.returncode)

        if result.returncode != 0:
            stderr = (result.stderr or "").strip()
            raise RuntimeError(
                f"LibreOffice exited with code {result.returncode}: {stderr}"
            )

        # Find the generated PDF
        pdf_files = list(Path(out_dir).glob("*.pdf"))
        if not pdf_files:
            raise RuntimeError(
                "LibreOffice ran successfully but produced no PDF file. "
                f"stdout={result.stdout!r} stderr={result.stderr!r}"
            )

        return pdf_files[0].read_bytes()

    except subprocess.TimeoutExpired:
        raise RuntimeError(
            f"LibreOffice conversion timed out after {timeout} seconds. "
            "Try a smaller file or contact support."
        )
    finally:
        shutil.rmtree(out_dir, ignore_errors=True)
        shutil.rmtree(profile, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════════════════════
# FALLBACK ENGINE — ReportLab  (pure Python, no system binary)
# ═══════════════════════════════════════════════════════════════════════════════

# ── Font registration ─────────────────────────────────────────────────────────

_F_REG  = "Helvetica"
_F_BOLD = "Helvetica-Bold"
_F_ITAL = "Helvetica-Oblique"
_F_BI   = "Helvetica-BoldOblique"

_FONT_CANDIDATES = [
    (
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Italic.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-BoldItalic.ttf",
    ),
    (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf",
    ),
    (
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Italic.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-BoldItalic.ttf",
    ),
]


def _register_fonts() -> None:
    global _F_REG, _F_BOLD, _F_ITAL, _F_BI
    for reg, bold, ital, bi in _FONT_CANDIDATES:
        if not os.path.exists(reg):
            continue
        try:
            pdfmetrics.registerFont(TTFont("_S1_REG",  reg))
            _F_REG = "_S1_REG"
            if os.path.exists(bold):
                pdfmetrics.registerFont(TTFont("_S1_BOLD", bold))
                _F_BOLD = "_S1_BOLD"
            if os.path.exists(ital):
                pdfmetrics.registerFont(TTFont("_S1_ITAL", ital))
                _F_ITAL = "_S1_ITAL"
            if os.path.exists(bi):
                pdfmetrics.registerFont(TTFont("_S1_BI",   bi))
                _F_BI   = "_S1_BI"
            logger.info("ReportLab: registered font family from %s", reg)
            return
        except Exception as exc:
            logger.debug("Font registration failed for %s: %s", reg, exc)


_register_fonts()

# ── Alignment map ─────────────────────────────────────────────────────────────

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT:    TA_LEFT,
    WD_ALIGN_PARAGRAPH.CENTER:  TA_CENTER,
    WD_ALIGN_PARAGRAPH.RIGHT:   TA_RIGHT,
    WD_ALIGN_PARAGRAPH.JUSTIFY: TA_JUSTIFY,
    None:                       TA_LEFT,
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def _emu_to_pt(emu: int) -> float:
    return emu / 12700.0


def _safe_pt(val, default: float = 11.0) -> float:
    if val is None:
        return default
    try:
        return float(val.pt)
    except Exception:
        return default


def _safe_color(run) -> Optional[str]:
    try:
        c = run.font.color
        if c and c.type is not None:
            rgb: RGBColor = c.rgb
            return "#{:02x}{:02x}{:02x}".format(rgb.red, rgb.green, rgb.blue)
    except Exception:
        pass
    return None


def _para_base_size(para, fallback: float = 11.0) -> float:
    try:
        if para.style and para.style.font and para.style.font.size:
            return _safe_pt(para.style.font.size, fallback)
    except Exception:
        pass
    return fallback


def _is_explicit_page_break(para) -> bool:
    for br in para._element.findall(".//" + qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _page_geometry(doc: Document):
    default_page   = A4
    default_margin = 2.5 * cm
    try:
        section = doc.sections[0]
        w_emu = section.page_width
        h_emu = section.page_height
        if w_emu and h_emu and w_emu > 0 and h_emu > 0:
            w_pt = _emu_to_pt(w_emu)
            h_pt = _emu_to_pt(h_emu)
            if abs(w_pt - 612) < 12 and abs(h_pt - 792) < 12:
                pagesize = LETTER
            elif abs(w_pt - 595) < 12 and abs(h_pt - 842) < 12:
                pagesize = A4
            else:
                pagesize = (w_pt, h_pt)
        else:
            pagesize = default_page

        def _m(emu, default):
            try:
                return _emu_to_pt(int(emu)) if emu else default
            except Exception:
                return default

        left   = _m(section.left_margin,   default_margin)
        right  = _m(section.right_margin,  default_margin)
        top    = _m(section.top_margin,    default_margin)
        bottom = _m(section.bottom_margin, default_margin)
        return pagesize, left, right, top, bottom
    except Exception:
        return default_page, default_margin, default_margin, default_margin, default_margin


class _NumCounters:
    def __init__(self) -> None:
        self._counts: dict = {}

    def next(self, num_id: str, ilvl: int) -> int:
        key = (num_id, ilvl)
        self._counts[key] = self._counts.get(key, 0) + 1
        return self._counts[key]

    def reset_below(self, num_id: str, ilvl: int) -> None:
        for k in list(self._counts.keys()):
            if k[0] == num_id and k[1] > ilvl:
                del self._counts[k]


def _get_list_info(para) -> tuple:
    style_name = (para.style.name or "").lower() if para.style else ""
    try:
        pPr   = para._element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
        if numPr is not None:
            numId_el = numPr.find(qn("w:numId"))
            ilvl_el  = numPr.find(qn("w:ilvl"))
            num_id   = numId_el.get(qn("w:val"), "0") if numId_el is not None else "0"
            ilvl     = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
            indent   = max(18.0, (ilvl + 1) * 18.0)
            is_bullet = "list bullet" in style_name or "bullet" in style_name
            return (is_bullet, num_id, ilvl, indent)
    except Exception:
        pass
    if "list bullet" in style_name:
        return (True, "0", 0, 18.0)
    if "list number" in style_name:
        return (False, "0", 0, 18.0)
    return (None, "0", 0, 0.0)


def _run_markup(para, base_size: float = 11.0) -> str:
    parts: list = []
    for run in para.runs:
        text = run.text or ""
        if not text:
            continue
        text = (
            escape(text, {'"': "&quot;", "'": "&apos;"})
            .replace("\r\n", "\n").replace("\r", "\n").replace("\n", "<br/>")
        )
        size  = _safe_pt(run.font.size, base_size)
        color = _safe_color(run)
        v_align = None
        try:
            v_align = run.font.vertAlign
        except Exception:
            pass

        if run.bold and run.italic:
            fname = _F_BI
        elif run.bold:
            fname = _F_BOLD
        elif run.italic:
            fname = _F_ITAL
        else:
            fname = _F_REG

        open_t:  list = [f'<font name="{fname}" size="{size:.1f}"']
        close_t: list = ["</font>"]

        if color:
            open_t[0] += f' color="{color}"'

        if v_align == "superscript":
            sup_size = max(6.0, size * 0.65)
            open_t[0] = f'<font name="{fname}" size="{sup_size:.1f}"'
            if color:
                open_t[0] += f' color="{color}"'
            open_t[0] += ">"
            open_t.append("<super>")
            close_t.insert(0, "</super>")
        elif v_align == "subscript":
            sub_size = max(6.0, size * 0.65)
            open_t[0] = f'<font name="{fname}" size="{sub_size:.1f}"'
            if color:
                open_t[0] += f' color="{color}"'
            open_t[0] += ">"
            open_t.append("<sub>")
            close_t.insert(0, "</sub>")
        else:
            open_t[0] += ">"

        if run.bold:
            open_t.append("<b>"); close_t.insert(0, "</b>")
        if run.italic:
            open_t.append("<i>"); close_t.insert(0, "</i>")
        if run.underline:
            open_t.append("<u>"); close_t.insert(0, "</u>")
        try:
            if run.font.strike:
                open_t.append("<strike>"); close_t.insert(0, "</strike>")
        except Exception:
            pass

        parts.append("".join(open_t) + text + "".join(close_t))
    return "".join(parts)


def _safe_paragraph(markup: str, style: ParagraphStyle) -> Paragraph:
    try:
        return Paragraph(markup, style)
    except Exception:
        plain = re.sub(r"<[^>]+>", "", markup)
        plain = plain.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
        try:
            return Paragraph(escape(plain), style)
        except Exception:
            return Paragraph("", style)


_style_seq = 0


def _make_style(
    size: float,
    leading: float,
    align: int,
    bold: bool = False,
    left_indent: float = 0,
    space_before: float = 0,
    space_after: float = 4,
) -> ParagraphStyle:
    global _style_seq
    _style_seq += 1
    return ParagraphStyle(
        name=f"_s1_{_style_seq}",
        fontName=_F_BOLD if bold else _F_REG,
        fontSize=size,
        leading=leading,
        alignment=align,
        leftIndent=left_indent,
        spaceBefore=space_before,
        spaceAfter=space_after,
    )


def _para_images(para, doc: Document, max_width_pt: float) -> list:
    images = []
    try:
        ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        for run in para.runs:
            for blip in run._element.findall(".//" + qn("a:blip"), namespaces=ns):
                rId = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if rId and rId in doc.part.rels:
                    img_part = doc.part.rels[rId].target_part
                    buf = io.BytesIO(img_part.blob)
                    try:
                        rl_img = RLImage(buf)
                        if rl_img.drawWidth > max_width_pt:
                            ratio = max_width_pt / rl_img.drawWidth
                            rl_img.drawWidth  *= ratio
                            rl_img.drawHeight *= ratio
                        images.append(rl_img)
                    except Exception:
                        pass
    except Exception:
        pass
    return images


def _reportlab_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Fallback: convert DOCX→PDF with ReportLab (pure Python)."""
    doc      = Document(io.BytesIO(docx_bytes))
    counters = _NumCounters()

    pagesize, left_m, right_m, top_m, bottom_m = _page_geometry(doc)
    page_w_pt = pagesize[0]
    text_w_pt = page_w_pt - left_m - right_m

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=pagesize,
        leftMargin=left_m,
        rightMargin=right_m,
        topMargin=top_m,
        bottomMargin=bottom_m,
    )

    base_style = _make_style(11, 14, TA_LEFT)
    story: list = []

    for para in doc.paragraphs:
        if _is_explicit_page_break(para):
            story.append(PageBreak())
            continue

        for img in _para_images(para, doc, text_w_pt):
            story.append(img)
            story.append(Spacer(1, 4))

        markup = _run_markup(para)
        if not markup.strip():
            story.append(Spacer(1, 6))
            continue

        align       = _ALIGN_MAP.get(para.alignment, TA_LEFT)
        style_lower = (para.style.name or "").lower() if para.style else ""

        if "title" in style_lower:
            ps = _make_style(22, 28, TA_CENTER, bold=True, space_before=0, space_after=12)
        elif "heading 1" in style_lower:
            ps = _make_style(18, 22, align, bold=True, space_before=10, space_after=8)
        elif "heading 2" in style_lower:
            ps = _make_style(15, 19, align, bold=True, space_before=8, space_after=6)
        elif "heading 3" in style_lower:
            ps = _make_style(13, 17, align, bold=True, space_before=6, space_after=4)
        elif "heading 4" in style_lower or "heading 5" in style_lower:
            ps = _make_style(12, 15, align, bold=True, space_before=4, space_after=3)
        else:
            is_bullet, num_id, ilvl, indent_pt = _get_list_info(para)
            if is_bullet is True:
                label  = "•" if ilvl == 0 else ("◦" if ilvl == 1 else "▪")
                markup = f"{label}&nbsp;&nbsp;{markup}"
                ps     = _make_style(11, 14, TA_LEFT, left_indent=indent_pt, space_after=2)
            elif is_bullet is False and num_id != "0":
                counters.reset_below(num_id, ilvl)
                n      = counters.next(num_id, ilvl)
                markup = f"{n}.&nbsp;&nbsp;{markup}"
                ps     = _make_style(11, 14, TA_LEFT, left_indent=indent_pt, space_after=2)
            else:
                sz = _para_base_size(para)
                ps = _make_style(sz, sz * 1.3, align, left_indent=indent_pt)

        story.append(_safe_paragraph(markup, ps))

    for table in doc.tables:
        rows: list = []
        for row in table.rows:
            cells: list = []
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                safe = escape(cell_text, {'"': "&quot;", "'": "&apos;"})
                cells.append(_safe_paragraph(safe, base_style))
            rows.append(cells)

        if not rows:
            continue
        col_n = max(len(r) for r in rows)
        if col_n == 0:
            continue
        for r in rows:
            while len(r) < col_n:
                r.append(_safe_paragraph("", base_style))

        col_w = text_w_pt / col_n
        tbl = Table(rows, colWidths=[col_w] * col_n)
        tbl.setStyle(TableStyle([
            ("GRID",          (0, 0), (-1, -1), 0.5, rl_colors.grey),
            ("BACKGROUND",    (0, 0), (-1,  0), rl_colors.HexColor("#DBEAFE")),
            ("FONTSIZE",      (0, 0), (-1, -1), 9),
            ("VALIGN",        (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING",   (0, 0), (-1, -1), 4),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
        ]))
        story.extend([Spacer(1, 8), tbl, Spacer(1, 8)])

    if not story:
        story.append(Paragraph("(empty document)", base_style))

    pdf.build(story)
    buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN CONVERTER — dispatches to LibreOffice or ReportLab
# ═══════════════════════════════════════════════════════════════════════════════

def convert_to_pdf(file_bytes: bytes, original_filename: str) -> tuple[bytes, str]:
    """
    Convert Word/ODT/RTF bytes to PDF bytes.

    Returns (pdf_bytes, engine_used).
    Raises RuntimeError on failure.
    """
    suffix = Path(secure_filename(original_filename)).suffix.lower() or ".docx"

    # ── LibreOffice path (primary) ───────────────────────────────────────────
    if _libreoffice_available():
        tmp_input = tempfile.NamedTemporaryFile(
            suffix=suffix, delete=False, prefix="lo_input_"
        )
        try:
            tmp_input.write(file_bytes)
            tmp_input.flush()
            tmp_input.close()
            pdf_bytes = _convert_with_libreoffice(tmp_input.name)
            return pdf_bytes, "libreoffice"
        finally:
            try:
                os.unlink(tmp_input.name)
            except OSError:
                pass

    # ── ReportLab fallback (DOCX only) ───────────────────────────────────────
    logger.warning("LibreOffice not found — falling back to ReportLab (DOCX only)")
    if suffix != ".docx":
        raise RuntimeError(
            f"LibreOffice is not available and ReportLab only supports .docx "
            f"(got {suffix}). Please install LibreOffice on the server."
        )
    pdf_bytes = _reportlab_docx_to_pdf(file_bytes)
    return pdf_bytes, "reportlab"


# ═══════════════════════════════════════════════════════════════════════════════
# Flask routes
# ═══════════════════════════════════════════════════════════════════════════════

@app.get("/")
def home():
    engine = "libreoffice" if _libreoffice_available() else "reportlab-fallback"
    return jsonify({
        "status":   "running",
        "tool":     "word-to-pdf",
        "engine":   engine,
        "accepts":  sorted(ACCEPTED_EXTENSIONS),
        "max_mb":   app.config["MAX_CONTENT_LENGTH"] // (1024 * 1024),
    })


@app.get("/health")
def health():
    lo_ok  = _libreoffice_available()
    status = "ok" if lo_ok else "degraded"
    return jsonify({
        "status":          status,
        "libreoffice":     lo_ok,
        "libreoffice_bin": _SOFFICE_BIN,
    }), 200


@app.route("/convert", methods=["POST"])
def convert_word():
    # ── Validate upload ───────────────────────────────────────────────────────
    upload = request.files.get("file")
    if not upload:
        return jsonify(error="No file uploaded."), 400
    if not upload.filename:
        return jsonify(error="No filename provided."), 400

    filename = secure_filename(upload.filename)
    suffix   = Path(filename).suffix.lower()

    if suffix not in ACCEPTED_EXTENSIONS:
        return jsonify(
            error=f"Unsupported file type '{suffix}'. "
                  f"Accepted: {', '.join(sorted(ACCEPTED_EXTENSIONS))}"
        ), 415

    file_bytes = upload.read()
    if not file_bytes:
        return jsonify(error="Uploaded file is empty."), 400

    stem = Path(filename).stem or "converted"

    # ── Convert ───────────────────────────────────────────────────────────────
    try:
        pdf_bytes, engine = convert_to_pdf(file_bytes, filename)
    except RuntimeError as exc:
        logger.error("Conversion failed: %s", exc)
        return jsonify(error=str(exc)), 500
    except Exception:
        logger.exception("Unexpected error during conversion")
        return jsonify(error="Internal conversion error. Please try again."), 500

    # ── Stream response ───────────────────────────────────────────────────────
    buf  = io.BytesIO(pdf_bytes)
    resp = send_file(
        buf,
        as_attachment=True,
        download_name=f"{stem}.pdf",
        mimetype="application/pdf",
        max_age=0,
    )
    resp.headers["Cache-Control"]     = "no-store"
    resp.headers["X-Conversion-Mode"] = engine
    resp.headers["X-Engine"]          = engine
    return resp


@app.errorhandler(413)
def file_too_large(_):
    max_mb = app.config["MAX_CONTENT_LENGTH"] // (1024 * 1024)
    return jsonify(error=f"File too large (max {max_mb} MB)."), 413


@app.errorhandler(415)
def unsupported_media(_):
    return jsonify(
        error=f"Unsupported media type. Accepted: {', '.join(sorted(ACCEPTED_EXTENSIONS))}"
    ), 415


# ── Local dev launcher ────────────────────────────────────────────────────────

if __name__ == "__main__":
    port  = int(os.getenv("PORT", "5000"))
    debug = os.getenv("FLASK_DEBUG", "0") == "1"
    logger.info("Starting word-to-pdf server on port %d (LibreOffice: %s)", port, _libreoffice_available())
    app.run(host="0.0.0.0", port=port, debug=debug)
