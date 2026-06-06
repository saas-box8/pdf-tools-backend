import gc
import io
import os
import re
import shutil
import tempfile
import traceback
from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from statistics import median
from typing import List, Tuple, Optional, Dict, Any

import fitz
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_UNDERLINE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from pdf2docx import Converter
from waitress import serve
from werkzeug.utils import secure_filename

# Optional RapidOCR (pure Python, no Tesseract/Poppler needed)
try:
    from rapidocr_onnxruntime import RapidOCR as _RapidOCR
    _ocr_engine = _RapidOCR()
except Exception:
    _RapidOCR = None
    _ocr_engine = None


# Optional RapidOCR (pure Python, no Tesseract/Poppler needed)
try:
    from rapidocr_onnxruntime import RapidOCR as _RapidOCR
    _ocr_engine = _RapidOCR()
except Exception:
    _RapidOCR = None
    _ocr_engine = None


# =========================================================
# APP CONFIGURATION
# =========================================================

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
app.config["MAX_CONTENT_LENGTH"] = 1024 * 1024 * 1024


# =========================================================
# FONT SUBSTITUTION MAP
# PDF font names → best Windows/web-safe equivalents
# =========================================================

FONT_MAP = {
    # Times / Serif
    "timesnewroman": "Times New Roman",
    "times": "Times New Roman",
    "timesroman": "Times New Roman",
    "timesnewromanps": "Times New Roman",
    "timesnewromanpsmt": "Times New Roman",
    # Arial / Helvetica
    "arial": "Arial",
    "helvetica": "Arial",
    "helveticaneue": "Arial",
    "arialmt": "Arial",
    "arialnarro": "Arial Narrow",
    # Calibri / Cambria
    "calibri": "Calibri",
    "cambria": "Cambria",
    "cambriamt": "Cambria",
    # Courier / Monospace
    "couriernew": "Courier New",
    "courier": "Courier New",
    "courierps": "Courier New",
    # Georgia
    "georgia": "Georgia",
    # Verdana
    "verdana": "Verdana",
    # Garamond
    "garamond": "Garamond",
    # Book Antiqua / Palatino
    "bookantiqua": "Book Antiqua",
    "palatino": "Palatino Linotype",
    "palatinolinotype": "Palatino Linotype",
    # Trebuchet
    "trebuchetms": "Trebuchet MS",
    # Tahoma
    "tahoma": "Tahoma",
    # Impact
    "impact": "Impact",
    # Myriad
    "myriadpro": "Arial",
    "myriad": "Arial",
    # Source Sans / Roboto (common in PDFs)
    "sourcesanspro": "Arial",
    "sourcesans": "Arial",
    "robotocondensed": "Arial Narrow",
    "roboto": "Arial",
    # Lato / Open Sans
    "lato": "Arial",
    "opensans": "Arial",
}


def resolve_font(raw_font: str) -> str:
    """Map a PDF font name to the best Windows-compatible font."""
    if not raw_font:
        return "Arial"
    # Strip subset prefix like "ABCDEF+"
    clean = raw_font.split("+")[-1].split(",")[0].strip()
    key = re.sub(r"[-_ ]", "", clean).lower()
    # Remove style suffixes for lookup
    for suffix in ("bold", "italic", "bolditalic", "medium", "regular", "light", "semibold", "black"):
        if key.endswith(suffix):
            key = key[: -len(suffix)]
            break
    return FONT_MAP.get(key, clean if clean else "Arial")


# =========================================================
# DATA CLASSES
# =========================================================

@dataclass
class TextSpan:
    text: str
    font: str = "Arial"
    size: float = 11.0
    color: Tuple[int, int, int] = (0, 0, 0)
    bg_color: Optional[Tuple[int, int, int]] = None
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strikethrough: bool = False
    superscript: bool = False
    subscript: bool = False
    bbox: Tuple[float, float, float, float] = (0, 0, 0, 0)
    origin: Tuple[float, float] = (0, 0)
    flags: int = 0
    char_spacing: float = 0.0


@dataclass
class TextLine:
    spans: List[TextSpan] = field(default_factory=list)
    bbox: Tuple[float, float, float, float] = (0, 0, 0, 0)
    y_position: float = 0.0
    height: float = 0.0
    width: float = 0.0
    page_width: float = 612.0
    is_empty: bool = False

    @property
    def full_text(self) -> str:
        return "".join(s.text for s in self.spans)

    @property
    def is_rtl(self) -> bool:
        text = self.full_text
        arabic_count = sum(1 for c in text if "\u0600" <= c <= "\u06FF" or "\u0590" <= c <= "\u05FF")
        return arabic_count > len(text) * 0.3

    @property
    def dominant_font(self) -> str:
        if not self.spans:
            return "Arial"
        counts: Dict[str, int] = defaultdict(int)
        for s in self.spans:
            counts[s.font] += len(s.text)
        return max(counts, key=counts.get)

    @property
    def dominant_size(self) -> float:
        if not self.spans:
            return 11.0
        counts: Dict[float, int] = defaultdict(int)
        for s in self.spans:
            counts[s.size] += len(s.text)
        return max(counts, key=counts.get)

    @property
    def is_bold(self) -> bool:
        if not self.spans:
            return False
        bold_chars = sum(len(s.text) for s in self.spans if s.bold)
        total_chars = sum(len(s.text) for s in self.spans)
        return bold_chars > total_chars * 0.5 if total_chars else False

    @property
    def is_italic(self) -> bool:
        if not self.spans:
            return False
        italic_chars = sum(len(s.text) for s in self.spans if s.italic)
        total_chars = sum(len(s.text) for s in self.spans)
        return italic_chars > total_chars * 0.5 if total_chars else False

    def detect_alignment(self, left_margin: float, right_margin: float) -> str:
        """Detect alignment using actual page geometry."""
        if not self.bbox:
            return "left"
        x0, _, x1, _ = self.bbox
        content_width = self.page_width - left_margin - right_margin
        if content_width <= 0:
            return "left"

        text_width = x1 - x0
        left_gap = x0 - left_margin
        right_gap = (self.page_width - right_margin) - x1
        center_of_text = (x0 + x1) / 2
        center_of_page = self.page_width / 2

        # Full-width text → justify candidate
        if text_width > content_width * 0.85:
            return "justify"
        # Center-aligned: text center near page center
        if abs(center_of_text - center_of_page) < content_width * 0.08:
            return "center"
        # Right-aligned: large left gap, small right gap
        if left_gap > content_width * 0.5 and right_gap < content_width * 0.15:
            return "right"
        return "left"


@dataclass
class TextParagraph:
    lines: List[TextLine] = field(default_factory=list)
    is_heading: bool = False
    heading_level: int = 0
    is_list_item: bool = False
    list_type: str = ""       # "bullet", "number", "letter", "roman"
    list_marker: str = ""
    list_level: int = 0
    is_code_block: bool = False
    space_before: float = 0.0
    space_after: float = 0.0
    left_indent: float = 0.0
    first_line_indent: float = 0.0
    alignment: str = "left"
    column: int = 0           # For multi-column layouts

    @property
    def full_text(self) -> str:
        return " ".join(line.full_text for line in self.lines)

    @property
    def bbox(self) -> Tuple[float, float, float, float]:
        if not self.lines:
            return (0, 0, 0, 0)
        x0 = min(ln.bbox[0] for ln in self.lines)
        y0 = min(ln.bbox[1] for ln in self.lines)
        x1 = max(ln.bbox[2] for ln in self.lines)
        y1 = max(ln.bbox[3] for ln in self.lines)
        return (x0, y0, x1, y1)


@dataclass
class TableData:
    rows: List[List[str]] = field(default_factory=list)
    bbox: Tuple[float, float, float, float] = (0, 0, 0, 0)
    has_header: bool = False
    cell_formats: List[List[Dict]] = field(default_factory=list)
    col_widths: List[float] = field(default_factory=list)
    header_bg: str = "D9E1F2"


@dataclass
class ImageData:
    image_bytes: bytes = b""
    width_inches: float = 0.0
    height_inches: float = 0.0
    bbox: Tuple[float, float, float, float] = (0, 0, 0, 0)
    page: int = 0
    caption: str = ""


@dataclass
class LinkData:
    text: str = ""
    url: str = ""
    bbox: Tuple[float, float, float, float] = (0, 0, 0, 0)
    page: int = 0


@dataclass
class HeaderFooterData:
    text_lines: List[str] = field(default_factory=list)
    is_header: bool = True
    page_number: Optional[str] = None


@dataclass
class PageData:
    page_num: int = 0
    width: float = 0.0
    height: float = 0.0
    rotation: int = 0
    left_margin: float = 72.0
    right_margin: float = 72.0
    top_margin: float = 72.0
    bottom_margin: float = 72.0
    paragraphs: List[TextParagraph] = field(default_factory=list)
    tables: List[TableData] = field(default_factory=list)
    images: List[ImageData] = field(default_factory=list)
    links: List[LinkData] = field(default_factory=list)
    header: Optional[HeaderFooterData] = None
    footer: Optional[HeaderFooterData] = None
    columns: int = 1          # Number of detected columns


# =========================================================
# UTILITY FUNCTIONS
# =========================================================

def safe_cleanup(path):
    try:
        if path and os.path.exists(path):
            shutil.rmtree(path, ignore_errors=True)
    except Exception:
        pass


def validate_pdf(file):
    if not file:
        raise Exception("No file uploaded")
    if file.filename == "":
        raise Exception("Empty filename")
    if not file.filename.lower().endswith(".pdf"):
        raise Exception("Only PDF files are allowed")


def extract_color(raw) -> Tuple[int, int, int]:
    """
    Extract RGB tuple from PyMuPDF color value.
    PyMuPDF can return:
      - int (packed sRGB: 0xRRGGBB)
      - float 0-1 (grayscale)
      - tuple of 1, 3, or 4 floats (gray, RGB, CMYK in 0-1 range)
    """
    if raw is None:
        return (0, 0, 0)
    if isinstance(raw, (int,)):
        r = (raw >> 16) & 0xFF
        g = (raw >> 8) & 0xFF
        b = raw & 0xFF
        return (r, g, b)
    if isinstance(raw, float):
        v = int(raw * 255)
        return (v, v, v)
    if isinstance(raw, (list, tuple)):
        if len(raw) == 1:
            v = int(raw[0] * 255)
            return (v, v, v)
        if len(raw) == 3:
            return (int(raw[0] * 255), int(raw[1] * 255), int(raw[2] * 255))
        if len(raw) == 4:
            # CMYK → RGB
            c, m, y, k = raw
            r = int(255 * (1 - c) * (1 - k))
            g = int(255 * (1 - m) * (1 - k))
            b = int(255 * (1 - y) * (1 - k))
            return (r, g, b)
    return (0, 0, 0)


def is_list_marker(text: str) -> Tuple[bool, str, str, int]:
    """Return (is_list, type, marker, indent_level)."""
    stripped = text.lstrip()
    leading_spaces = len(text) - len(stripped)
    level = leading_spaces // 3  # Approximate nesting

    bullet_re = re.compile(
        r"^([•∙○●◦▪▫►▸▹◆◇★☆✦✧→←–—\-\*\+])\s+"
    )
    m = bullet_re.match(stripped)
    if m:
        return True, "bullet", m.group(), level

    num_patterns = [
        (r"^(\d{1,3})\.\s+", "number"),
        (r"^(\d{1,3})\)\s+", "number"),
        (r"^\((\d{1,3})\)\s+", "number"),
        (r"^([ivxlcdmIVXLCDM]{1,6})\.\s+", "roman"),
        (r"^([a-zA-Z])\.\s+", "letter"),
        (r"^([a-zA-Z])\)\s+", "letter"),
        (r"^\(([a-zA-Z])\)\s+", "letter"),
    ]
    for pattern, list_type in num_patterns:
        m = re.match(pattern, stripped)
        if m:
            return True, list_type, m.group(), level

    return False, "", "", 0


def is_page_number(text: str) -> bool:
    t = text.strip()
    return bool(re.fullmatch(r"[\-–—]?\s*\d+\s*[\-–—]?", t)) or t.lower() in ("page", "")


def get_alignment_value(alignment: str):
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)


def set_cell_shading(cell, color_hex: str):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text: str, url: str):
    try:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = parse_xml(
            f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            f' r:id="{r_id}">'
            f'<w:r><w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
            f"<w:t>{text}</w:t></w:r></w:hyperlink>"
        )
        paragraph._p.append(hyperlink)
    except Exception as e:
        paragraph.add_run(text)


# =========================================================
# DOCUMENT-LEVEL FONT SIZE ANALYSIS
# =========================================================

class FontSizeAnalyzer:
    """Analyze font sizes across the whole document to determine body size and heading thresholds."""

    def __init__(self):
        self.sizes: List[float] = []

    def feed_page(self, page):
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        self.sizes.append(span.get("size", 11))

    def body_size(self) -> float:
        """Most common (modal) font size = body text size."""
        if not self.sizes:
            return 11.0
        counts: Dict[float, int] = defaultdict(int)
        for s in self.sizes:
            rounded = round(s * 2) / 2  # Round to nearest 0.5pt
            counts[rounded] += 1
        return max(counts, key=counts.get)

    def is_heading(self, size: float, is_bold: bool, text: str, body_size: float) -> Tuple[bool, int]:
        """
        Determine heading level relative to document's body font size.
        This is far more accurate than absolute thresholds.
        """
        text = text.strip()
        if not text or len(text) > 200:
            return False, 0

        ratio = size / body_size if body_size > 0 else 1.0

        # Ends with sentence punctuation → unlikely heading
        if text.endswith((".", "?", "!", ",", ";")):
            if ratio < 1.4:
                return False, 0

        if ratio >= 2.0:
            return True, 1
        if ratio >= 1.6:
            return True, 2
        if ratio >= 1.35:
            return True, 3
        if ratio >= 1.15 and is_bold:
            return True, 4
        if ratio >= 1.0 and is_bold and len(text) < 80:
            # Check for ALL-CAPS headings at body size
            alpha = [c for c in text if c.isalpha()]
            if alpha and sum(1 for c in alpha if c.isupper()) / len(alpha) > 0.8:
                return True, 3
            return True, 4

        return False, 0


# =========================================================
# MULTI-COLUMN DETECTOR
# =========================================================

def detect_columns(lines: list, page_width: float) -> int:
    """
    Detect number of text columns by analyzing X-start distribution.
    Returns 1 (single column) or 2+ (multi-column).
    """
    if not lines:
        return 1
    x_starts = [ln.bbox[0] for ln in lines if not ln.is_empty]
    if len(x_starts) < 6:
        return 1

    # Find the median gap between column starts
    x_starts.sort()
    half = page_width / 2

    # Count lines on each half
    left_count = sum(1 for x in x_starts if x < half * 0.6)
    right_count = sum(1 for x in x_starts if x > half * 0.8)

    if left_count > 3 and right_count > 3:
        # Check if right column start is consistently different from left
        left_starts = [x for x in x_starts if x < half * 0.6]
        right_starts = [x for x in x_starts if x > half * 0.8]
        if left_starts and right_starts:
            left_med = median(left_starts)
            right_med = median(right_starts)
            if right_med - left_med > page_width * 0.3:
                return 2
    return 1


# =========================================================
# ADVANCED PDF EXTRACTOR
# =========================================================

class AdvancedPDFExtractor:
    """High-fidelity PDF content extractor using PyMuPDF."""

    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.doc = fitz.open(pdf_path)
        self.pages_data: List[PageData] = []
        self._font_analyzer = FontSizeAnalyzer()
        self._body_size: float = 11.0
        self._para_gap_threshold: float = 8.0

    def extract_all(self) -> List[PageData]:
        # Pass 1: Analyze font sizes across entire document
        for page_num in range(len(self.doc)):
            self._font_analyzer.feed_page(self.doc[page_num])
        self._body_size = self._font_analyzer.body_size()

        # Pass 2: Full extraction
        for page_num in range(len(self.doc)):
            page = self.doc[page_num]
            page_data = self._extract_page(page, page_num)
            self.pages_data.append(page_data)

        return self.pages_data

    # ----------------------------------------------------------
    # PAGE-LEVEL EXTRACTION
    # ----------------------------------------------------------

    def _extract_page(self, page, page_num: int) -> PageData:
        page_data = PageData(
            page_num=page_num,
            width=page.rect.width,
            height=page.rect.height,
            rotation=page.rotation,
        )

        # Estimate margins from text bounding boxes
        self._estimate_margins(page, page_data)

        # Tables (extract first so we can mask their bboxes from text extraction)
        tables = self._extract_tables(page)
        table_bboxes = [t.bbox for t in tables]
        page_data.tables = tables

        # Images
        page_data.images = self._extract_images(page, page_num, page_data)

        # Hyperlinks
        page_data.links = self._extract_links(page, page_num)

        # Text
        paragraphs = self._extract_text_blocks(page, table_bboxes, page_data)
        page_data.paragraphs = paragraphs

        # Column detection
        all_lines = [ln for p in paragraphs for ln in p.lines]
        page_data.columns = detect_columns(all_lines, page_data.width)

        # Header / footer separation
        self._detect_header_footer(page_data)

        return page_data

    def _estimate_margins(self, page, page_data: PageData):
        """Estimate page margins by looking at outermost text bboxes."""
        text_dict = page.get_text("dict")
        x_mins, x_maxs, y_mins, y_maxs = [], [], [], []

        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            bbox = block.get("bbox", ())
            if len(bbox) == 4:
                x_mins.append(bbox[0])
                x_maxs.append(bbox[2])
                y_mins.append(bbox[1])
                y_maxs.append(bbox[3])

        pw, ph = page_data.width, page_data.height
        page_data.left_margin = min(x_mins) if x_mins else pw * 0.1
        page_data.right_margin = pw - max(x_maxs) if x_maxs else pw * 0.1
        page_data.top_margin = min(y_mins) if y_mins else ph * 0.1
        page_data.bottom_margin = ph - max(y_maxs) if y_maxs else ph * 0.1

        # Clamp to reasonable values (0.25" – 2")
        min_m, max_m = 18, 144
        page_data.left_margin = max(min_m, min(max_m, page_data.left_margin))
        page_data.right_margin = max(min_m, min(max_m, page_data.right_margin))
        page_data.top_margin = max(min_m, min(max_m, page_data.top_margin))
        page_data.bottom_margin = max(min_m, min(max_m, page_data.bottom_margin))

    # ----------------------------------------------------------
    # TABLE EXTRACTION
    # ----------------------------------------------------------

    def _extract_tables(self, page) -> List[TableData]:
        tables = []
        try:
            found = page.find_tables()
            for tbl in found:
                td = TableData(
                    bbox=tuple(tbl.bbox),
                    has_header=True,
                )
                raw = tbl.extract()
                if not raw:
                    continue

                # Determine column widths from header row bbox
                try:
                    header_cells = tbl.rows[0].cells
                    col_widths = [abs(c[2] - c[0]) for c in header_cells]
                    td.col_widths = col_widths
                except Exception:
                    td.col_widths = []

                for row_idx, row in enumerate(raw):
                    cleaned_row = [str(c).strip() if c is not None else "" for c in row]
                    td.rows.append(cleaned_row)
                    td.cell_formats.append(
                        [{"bold": row_idx == 0, "text": c} for c in cleaned_row]
                    )

                if td.rows:
                    tables.append(td)
        except Exception as e:
            print(f"[TABLE] {e}")
        return tables

    # ----------------------------------------------------------
    # IMAGE EXTRACTION
    # ----------------------------------------------------------

    def _extract_images(self, page, page_num: int, page_data: PageData) -> List[ImageData]:
        images = []
        dpi = 150
        seen_xrefs: set = set()

        try:
            for img_info in page.get_images(full=True):
                xref = img_info[0]
                if xref in seen_xrefs:
                    continue
                seen_xrefs.add(xref)

                try:
                    base_image = self.doc.extract_image(xref)
                    if not base_image or not base_image.get("image"):
                        continue

                    # Get actual position on page
                    rects = page.get_image_rects(xref)
                    bbox = (0, 0, 0, 0)
                    if rects:
                        r = rects[0]
                        bbox = (r.x0, r.y0, r.x1, r.y1)
                        # Use actual rendered dimensions from page rect
                        w_pts = r.x1 - r.x0
                        h_pts = r.y1 - r.y0
                        width_inches = w_pts / 72.0
                        height_inches = h_pts / 72.0
                    else:
                        px_w = base_image.get("width", 0)
                        px_h = base_image.get("height", 0)
                        width_inches = px_w / dpi if px_w else 4.0
                        height_inches = px_h / dpi if px_h else 3.0

                    # Cap at 6 inches wide
                    max_w = (page_data.width - page_data.left_margin - page_data.right_margin) / 72.0
                    if width_inches > max_w:
                        scale = max_w / width_inches
                        width_inches = max_w
                        height_inches *= scale

                    images.append(
                        ImageData(
                            image_bytes=base_image["image"],
                            width_inches=width_inches,
                            height_inches=height_inches,
                            bbox=bbox,
                            page=page_num,
                        )
                    )
                except Exception as e:
                    print(f"[IMAGE xref={xref}] {e}")
        except Exception as e:
            print(f"[IMAGES] {e}")

        return images

    # ----------------------------------------------------------
    # HYPERLINK EXTRACTION
    # ----------------------------------------------------------

    def _extract_links(self, page, page_num: int) -> List[LinkData]:
        links = []
        try:
            for link in page.get_links():
                if link.get("kind") == fitz.LINK_URI:
                    rect = fitz.Rect(link.get("from") or link.get("rect") or [0, 0, 0, 0])
                    text = page.get_text("text", clip=rect).strip()
                    url = link.get("uri", "")
                    if url:
                        links.append(
                            LinkData(
                                text=text or url,
                                url=url,
                                bbox=(rect.x0, rect.y0, rect.x1, rect.y1),
                                page=page_num,
                            )
                        )
        except Exception as e:
            print(f"[LINKS] {e}")
        return links

    # ----------------------------------------------------------
    # TEXT BLOCK EXTRACTION
    # ----------------------------------------------------------

    def _extract_text_blocks(
        self, page, exclude_bboxes: List[Tuple], page_data: PageData
    ) -> List[TextParagraph]:
        raw_dict = page.get_text("rawdict")  # rawdict gives char-level detail
        all_spans: List[TextSpan] = []

        # Build a URL map: bbox → url for hyperlink decoration
        link_map = {lnk.bbox: lnk.url for lnk in page_data.links}

        for block in raw_dict.get("blocks", []):
            if block.get("type") != 0:
                continue

            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text.strip():
                        continue

                    span_bbox = span.get("bbox", (0, 0, 0, 0))
                    # Skip text inside detected table areas
                    if self._is_in_any_bbox(span_bbox, exclude_bboxes):
                        continue

                    raw_color = span.get("color", 0)
                    color = extract_color(raw_color)

                    flags = span.get("flags", 0)
                    # Flag bits: 1=superscript, 2=italic, 4=serifed, 8=monospaced,
                    #            16=bold, 32=??? (varies by version), 64=strikethrough
                    bold = bool(flags & 16)
                    italic = bool(flags & 2)
                    underline = bool(flags & 4) or span.get("underline", False)
                    strikethrough = bool(flags & 64) or span.get("strikethrough", False)

                    # Superscript / subscript detection via origin Y offset
                    origin_y = span.get("origin", (0, 0))[1]
                    bbox_y_center = (span_bbox[1] + span_bbox[3]) / 2
                    line_y_top = line["bbox"][1]
                    line_y_bot = line["bbox"][3]
                    line_height = line_y_bot - line_y_top
                    superscript = subscript = False
                    if line_height > 0:
                        # Span bbox top is high on line → superscript
                        if span_bbox[1] < line_y_top + line_height * 0.3 and span_bbox[3] < line_y_bot - line_height * 0.2:
                            superscript = True
                        # Span bbox bottom is low on line → subscript
                        elif span_bbox[3] > line_y_bot - line_height * 0.3 and span_bbox[1] > line_y_top + line_height * 0.2:
                            subscript = True

                    raw_font = span.get("font", "Arial")
                    # Also detect bold/italic from font name if flags missed it
                    font_lower = raw_font.lower()
                    if not bold and ("bold" in font_lower or "black" in font_lower or "heavy" in font_lower):
                        bold = True
                    if not italic and ("italic" in font_lower or "oblique" in font_lower):
                        italic = True

                    resolved_font = resolve_font(raw_font)

                    ts = TextSpan(
                        text=text,
                        font=resolved_font,
                        size=span.get("size", 11),
                        color=color,
                        bold=bold,
                        italic=italic,
                        underline=underline,
                        strikethrough=strikethrough,
                        superscript=superscript,
                        subscript=subscript,
                        bbox=span_bbox,
                        origin=span.get("origin", (0, 0)),
                        flags=flags,
                    )
                    all_spans.append(ts)

        lines = self._group_spans_into_lines(all_spans, page_data.width)
        paragraphs = self._group_lines_into_paragraphs(lines, page_data)
        self._post_process_paragraphs(paragraphs, page_data)
        return paragraphs

    def _is_in_any_bbox(
        self, span_bbox: Tuple, exclude: List[Tuple], tolerance: float = 1.0
    ) -> bool:
        px0, py0, px1, py1 = span_bbox
        for bx0, by0, bx1, by1 in exclude:
            if (
                px0 >= bx0 - tolerance
                and py0 >= by0 - tolerance
                and px1 <= bx1 + tolerance
                and py1 <= by1 + tolerance
            ):
                return True
        return False

    # ----------------------------------------------------------
    # SPAN → LINE GROUPING
    # ----------------------------------------------------------

    def _group_spans_into_lines(
        self, spans: List[TextSpan], page_width: float
    ) -> List[TextLine]:
        if not spans:
            return []

        # Sort top-to-bottom, left-to-right
        spans = sorted(spans, key=lambda s: (round(s.origin[1], 0), s.origin[0]))

        lines: List[TextLine] = []
        current: List[TextSpan] = [spans[0]]
        current_y = spans[0].origin[1]
        current_h = spans[0].bbox[3] - spans[0].bbox[1]

        for span in spans[1:]:
            sy = span.origin[1]
            sh = span.bbox[3] - span.bbox[1]
            avg_h = (current_h + sh) / 2.0 if sh > 0 else current_h
            tol = max(3.0, avg_h * 0.35)

            if abs(sy - current_y) <= tol:
                current.append(span)
                if sh > current_h:
                    current_h = sh
            else:
                line = self._make_line(current, page_width)
                if not line.is_empty:
                    lines.append(line)
                current = [span]
                current_y = sy
                current_h = sh

        if current:
            line = self._make_line(current, page_width)
            if not line.is_empty:
                lines.append(line)

        return lines

    def _make_line(self, spans: List[TextSpan], page_width: float) -> TextLine:
        if not spans:
            return TextLine(is_empty=True)
        spans_sorted = sorted(spans, key=lambda s: s.bbox[0])
        x0 = min(s.bbox[0] for s in spans_sorted)
        y0 = min(s.bbox[1] for s in spans_sorted)
        x1 = max(s.bbox[2] for s in spans_sorted)
        y1 = max(s.bbox[3] for s in spans_sorted)
        return TextLine(
            spans=spans_sorted,
            bbox=(x0, y0, x1, y1),
            y_position=y0,
            height=y1 - y0,
            width=x1 - x0,
            page_width=page_width,
            is_empty=False,
        )

    # ----------------------------------------------------------
    # LINE → PARAGRAPH GROUPING
    # ----------------------------------------------------------

    def _group_lines_into_paragraphs(
        self, lines: List[TextLine], page_data: PageData
    ) -> List[TextParagraph]:
        if not lines:
            return []

        paragraphs: List[TextParagraph] = []
        current_lines: List[TextLine] = [lines[0]]
        prev = lines[0]

        for line in lines[1:]:
            gap = line.y_position - (prev.y_position + prev.height)
            avg_h = max(line.height, prev.height, 1.0)
            rel_gap = gap / avg_h

            same = self._is_same_paragraph(prev, line, gap, rel_gap)

            if same:
                current_lines.append(line)
            else:
                paragraphs.append(self._make_paragraph(current_lines, page_data))
                current_lines = [line]
            prev = line

        if current_lines:
            paragraphs.append(self._make_paragraph(current_lines, page_data))

        return paragraphs

    def _is_same_paragraph(
        self, prev: TextLine, curr: TextLine, gap: float, rel_gap: float
    ) -> bool:
        # Very large gap → definitely new paragraph
        if gap > self._para_gap_threshold * 2:
            return False

        # Different columns (large X jump) → new paragraph
        x_jump = abs(curr.bbox[0] - prev.bbox[0])
        if x_jump > prev.page_width * 0.3 and gap > 2:
            return False

        # Size change → likely new paragraph
        if abs(prev.dominant_size - curr.dominant_size) > 2.5:
            return False

        # Bold→non-bold or vice-versa → new paragraph
        if prev.is_bold != curr.is_bold:
            return False

        # Moderate gap with similar indent → same paragraph
        if rel_gap < 0.5 and x_jump < 20:
            return True

        # Indent suggests paragraph continuation
        # (right edge of prev line close to page width → text wrapped)
        content_width = prev.page_width - 72 * 2  # rough content width
        if prev.width > content_width * 0.7 and rel_gap < 1.2:
            return True

        # Small absolute gap
        if gap < self._para_gap_threshold and rel_gap < 0.8:
            return True

        return False

    def _make_paragraph(
        self, lines: List[TextLine], page_data: PageData
    ) -> TextParagraph:
        para = TextParagraph(lines=lines)
        if not lines:
            return para

        # Alignment (use the first full-width line or first line)
        para.alignment = lines[0].detect_alignment(
            page_data.left_margin, page_data.right_margin
        )

        # Left indent relative to page margin
        min_x = min(ln.bbox[0] for ln in lines)
        para.left_indent = max(0.0, min_x - page_data.left_margin)

        # First-line indent
        if len(lines) > 1:
            rest_x = min(ln.bbox[0] for ln in lines[1:])
            para.first_line_indent = lines[0].bbox[0] - rest_x

        return para

    # ----------------------------------------------------------
    # POST-PROCESSING
    # ----------------------------------------------------------

    def _post_process_paragraphs(
        self, paragraphs: List[TextParagraph], page_data: PageData
    ):
        for para in paragraphs:
            text = para.full_text.strip()
            if not text:
                continue

            # List detection
            is_list, list_type, marker, level = is_list_marker(text)
            if is_list:
                para.is_list_item = True
                para.list_type = list_type
                para.list_marker = marker
                para.list_level = level
                # Strip marker from first span
                if para.lines and para.lines[0].spans:
                    sp = para.lines[0].spans[0]
                    sp.text = sp.text.lstrip()
                    sp.text = re.sub(r"^" + re.escape(marker.strip()), "", sp.text).lstrip()

            # Heading detection
            if para.lines:
                fl = para.lines[0]
                is_h, hlevel = self._font_analyzer.is_heading(
                    fl.dominant_size, fl.is_bold, text, self._body_size
                )
                if is_h:
                    para.is_heading = True
                    para.heading_level = hlevel

            # Code block detection: monospaced font
            if para.lines:
                fl = para.lines[0]
                font_lower = fl.dominant_font.lower()
                if any(m in font_lower for m in ("courier", "mono", "consolas", "inconsolata", "sourcecodepro")):
                    para.is_code_block = True

    # ----------------------------------------------------------
    # HEADER / FOOTER DETECTION
    # ----------------------------------------------------------

    def _detect_header_footer(self, page_data: PageData):
        ph = page_data.height
        header_zone = ph * 0.08
        footer_zone = ph * 0.92

        main = []
        for para in page_data.paragraphs:
            bbox = para.bbox
            if not bbox or bbox == (0, 0, 0, 0):
                main.append(para)
                continue

            y0, y1 = bbox[1], bbox[3]

            if y1 < header_zone and not para.is_heading:
                if page_data.header is None:
                    page_data.header = HeaderFooterData(is_header=True)
                page_data.header.text_lines.append(para.full_text)
            elif y0 > footer_zone:
                if page_data.footer is None:
                    page_data.footer = HeaderFooterData(is_header=False)
                t = para.full_text.strip()
                if is_page_number(t):
                    page_data.footer.page_number = t
                else:
                    page_data.footer.text_lines.append(t)
            else:
                main.append(para)

        page_data.paragraphs = main

    def close(self):
        if self.doc:
            self.doc.close()


# =========================================================
# ADVANCED DOCX BUILDER
# =========================================================

class AdvancedDocxBuilder:
    """Build high-fidelity Word documents from extracted PDF data."""

    # List reference counters (per list type, so they restart correctly)
    _bullet_refs: Dict[str, int] = {}

    def __init__(self, pages_data: List[PageData]):
        self.pages_data = pages_data
        self.doc = Document()
        self._body_font = "Arial"
        self._body_size = 11.0
        self._numbering_idx = 0       # For unique list numbering references
        self._list_refs: Dict[str, str] = {}  # list_type → numbering abstract id
        self._setup_styles()
        self._setup_numbering()

    # ----------------------------------------------------------
    # STYLES & NUMBERING
    # ----------------------------------------------------------

    def _setup_styles(self):
        """Configure base document styles."""
        normal = self.doc.styles["Normal"]
        normal.font.name = self._body_font
        normal.font.size = Pt(self._body_size)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _setup_numbering(self):
        """Inject numbering XML into the document for bullet and ordered lists."""
        numbering_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            # Abstract num 0 → Bullets (up to 3 levels)
            '<w:abstractNum w:abstractNumId="0">'
            + "".join(
                f'<w:lvl w:ilvl="{i}">'
                f'<w:start w:val="1"/>'
                f'<w:numFmt w:val="bullet"/>'
                f'<w:lvlText w:val="&#x2022;"/>'
                f'<w:lvlJc w:val="left"/>'
                f'<w:pPr><w:ind w:left="{720 + i * 360}" w:hanging="360"/></w:pPr>'
                f'</w:lvl>'
                for i in range(9)
            )
            + "</w:abstractNum>"
            # Abstract num 1 → Decimal numbers
            '<w:abstractNum w:abstractNumId="1">'
            + "".join(
                f'<w:lvl w:ilvl="{i}">'
                f'<w:start w:val="1"/>'
                f'<w:numFmt w:val="decimal"/>'
                f'<w:lvlText w:val="%{i+1}."/>'
                f'<w:lvlJc w:val="left"/>'
                f'<w:pPr><w:ind w:left="{720 + i * 360}" w:hanging="360"/></w:pPr>'
                f'</w:lvl>'
                for i in range(9)
            )
            + "</w:abstractNum>"
            # Abstract num 2 → Letters
            '<w:abstractNum w:abstractNumId="2">'
            + "".join(
                f'<w:lvl w:ilvl="{i}">'
                f'<w:start w:val="1"/>'
                f'<w:numFmt w:val="lowerLetter"/>'
                f'<w:lvlText w:val="%{i+1}."/>'
                f'<w:lvlJc w:val="left"/>'
                f'<w:pPr><w:ind w:left="{720 + i * 360}" w:hanging="360"/></w:pPr>'
                f'</w:lvl>'
                for i in range(9)
            )
            + "</w:abstractNum>"
            # Abstract num 3 → Roman numerals
            '<w:abstractNum w:abstractNumId="3">'
            + "".join(
                f'<w:lvl w:ilvl="{i}">'
                f'<w:start w:val="1"/>'
                f'<w:numFmt w:val="lowerRoman"/>'
                f'<w:lvlText w:val="%{i+1}."/>'
                f'<w:lvlJc w:val="left"/>'
                f'<w:pPr><w:ind w:left="{720 + i * 360}" w:hanging="360"/></w:pPr>'
                f'</w:lvl>'
                for i in range(9)
            )
            + "</w:abstractNum>"
            # Concrete num instances
            '<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
            '<w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num>'
            '<w:num w:numId="3"><w:abstractNumId w:val="2"/></w:num>'
            '<w:num w:numId="4"><w:abstractNumId w:val="3"/></w:num>'
            "</w:numbering>"
        )

        # Replace existing numbering part or add one
        try:
            from docx.opc.part import Part
            from docx.opc.packuri import PackURI
            from docx.oxml import parse_xml as px

            numbering_part = self.doc.part.numbering_part
            numbering_part._element = px(numbering_xml.encode("utf-8"))
        except Exception:
            pass  # Numbering will fall back to manual markers

    # ----------------------------------------------------------
    # BUILD ENTRY POINT
    # ----------------------------------------------------------

    def build(self, output_path: str):
        if not self.pages_data:
            self.doc.save(output_path)
            return

        self._setup_section(self.pages_data[0])

        for idx, page_data in enumerate(self.pages_data):
            if idx > 0:
                self.doc.add_page_break()
                if (
                    page_data.width != self.pages_data[0].width
                    or page_data.height != self.pages_data[0].height
                ):
                    self._setup_section(page_data)

            self._add_headers_footers(page_data)
            self._add_page_content(page_data)

        self.doc.save(output_path)

    # ----------------------------------------------------------
    # SECTION SETUP
    # ----------------------------------------------------------

    def _setup_section(self, page_data: PageData):
        sec = self.doc.sections[-1]
        sec.page_width = Inches(page_data.width / 72.0)
        sec.page_height = Inches(page_data.height / 72.0)
        sec.top_margin = Inches(max(0.25, page_data.top_margin / 72.0))
        sec.bottom_margin = Inches(max(0.25, page_data.bottom_margin / 72.0))
        sec.left_margin = Inches(max(0.25, page_data.left_margin / 72.0))
        sec.right_margin = Inches(max(0.25, page_data.right_margin / 72.0))

    # ----------------------------------------------------------
    # HEADER / FOOTER
    # ----------------------------------------------------------

    def _add_headers_footers(self, page_data: PageData):
        sec = self.doc.sections[-1]

        if page_data.header:
            hdr = sec.header
            hdr.is_linked_to_previous = False
            p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            p.clear()
            for text in page_data.header.text_lines:
                run = p.add_run(text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(120, 120, 120)

        if page_data.footer:
            ftr = sec.footer
            ftr.is_linked_to_previous = False
            p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.clear()
            if page_data.footer.page_number:
                run = p.add_run(page_data.footer.page_number)
                run.font.size = Pt(9)
            for text in page_data.footer.text_lines:
                run = p.add_run(text)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(120, 120, 120)

    # ----------------------------------------------------------
    # PAGE CONTENT
    # ----------------------------------------------------------

    def _add_page_content(self, page_data: PageData):
        """Add all content sorted by Y position."""
        elements = []
        for para in page_data.paragraphs:
            elements.append(("p", para.bbox[1], para))
        for tbl in page_data.tables:
            elements.append(("t", tbl.bbox[1], tbl))
        for img in page_data.images:
            elements.append(("i", img.bbox[1], img))

        elements.sort(key=lambda e: e[1])

        for kind, _, data in elements:
            if kind == "p":
                self._add_paragraph(data, page_data)
            elif kind == "t":
                self._add_table(data, page_data)
            elif kind == "i":
                self._add_image(data)

    # ----------------------------------------------------------
    # PARAGRAPH
    # ----------------------------------------------------------

    def _add_paragraph(self, para: TextParagraph, page_data: PageData):
        if not para.lines:
            return

        text = para.full_text.strip()
        if not text:
            return

        # Heading
        if para.is_heading and not para.is_list_item:
            heading_style = {
                1: "Heading 1",
                2: "Heading 2",
                3: "Heading 3",
                4: "Heading 4",
            }.get(para.heading_level, "Heading 4")
            p = self.doc.add_paragraph(style=heading_style)
        else:
            p = self.doc.add_paragraph()

        # Alignment
        p.alignment = get_alignment_value(para.alignment)

        # Spacing
        pf = p.paragraph_format
        if para.space_before > 0:
            pf.space_before = Pt(min(para.space_before / 2.0, 24))
        if para.space_after > 0:
            pf.space_after = Pt(min(para.space_after / 2.0, 24))

        # Indentation (convert PDF points → inches)
        content_width = page_data.width - page_data.left_margin - page_data.right_margin
        if para.left_indent > 0:
            pf.left_indent = Inches(min(para.left_indent / 72.0, 3.0))
        if abs(para.first_line_indent) > 3:
            pf.first_line_indent = Inches(
                max(-0.5, min(0.5, para.first_line_indent / 72.0))
            )

        # RTL paragraph
        if para.lines and para.lines[0].is_rtl:
            pPr = p._p.get_or_add_pPr()
            bidi_el = parse_xml(
                f'<w:bidi {nsdecls("w")}/>'
            )
            pPr.append(bidi_el)

        # List formatting
        if para.is_list_item:
            self._apply_list_numbering(p, para)

        # Code block: monospace + light background
        if para.is_code_block:
            pf.left_indent = Inches(0.25)
            try:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F5F5"/>'
                )
                p._p.get_or_add_pPr().append(shading)
            except Exception:
                pass

        # Add runs
        for line_idx, line in enumerate(para.lines):
            if line_idx > 0:
                run = p.add_run()
                run.add_break()

            for span in line.spans:
                self._add_run(p, span)

    def _add_run(self, paragraph, span: TextSpan):
        run = paragraph.add_run(span.text)

        # Font name
        fn = span.font or self._body_font
        run.font.name = fn
        # Set cs (complex script) font too for RTL / CJK support
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} '
                f'w:ascii="{fn}" w:hAnsi="{fn}" '
                f'w:eastAsia="{fn}" w:cs="{fn}"/>'
            )
            rPr.append(rFonts)
        else:
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rFonts.set(qn(attr), fn)

        # Size
        if span.size and span.size > 0:
            run.font.size = Pt(span.size)

        # Color (skip near-black to let themes apply)
        r, g, b = span.color
        if not (r < 20 and g < 20 and b < 20):
            run.font.color.rgb = RGBColor(r, g, b)

        # Bold / Italic
        run.bold = span.bold
        run.italic = span.italic

        # Underline
        if span.underline:
            run.underline = True

        # Strikethrough
        if span.strikethrough:
            run.font.strike = True

        # Superscript / Subscript
        if span.superscript:
            run.font.superscript = True
        elif span.subscript:
            run.font.subscript = True

        # Background highlight
        if span.bg_color:
            br, bg, bb = span.bg_color
            # Only apply if not white/near-white
            if not (br > 240 and bg > 240 and bb > 240):
                color_hex = f"{br:02X}{bg:02X}{bb:02X}"
                try:
                    shd = parse_xml(
                        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
                    )
                    rPr.append(shd)
                except Exception:
                    pass

    def _apply_list_numbering(self, paragraph, para: TextParagraph):
        """Apply proper docx numbering to a list paragraph."""
        level = max(0, min(8, para.list_level))
        num_id_map = {
            "bullet": "1",
            "number": "2",
            "letter": "3",
            "roman": "4",
        }
        num_id = num_id_map.get(para.list_type, "1")

        try:
            pPr = paragraph._p.get_or_add_pPr()
            numPr = parse_xml(
                f'<w:numPr {nsdecls("w")}>'
                f'<w:ilvl w:val="{level}"/>'
                f'<w:numId w:val="{num_id}"/>'
                f"</w:numPr>"
            )
            # Remove any existing numPr
            existing = pPr.find(qn("w:numPr"))
            if existing is not None:
                pPr.remove(existing)
            pPr.insert(0, numPr)
        except Exception:
            # Fallback: manual bullet
            run = paragraph.add_run(
                ("• " if para.list_type == "bullet" else f"{para.list_marker} ")
            )
            run.font.name = self._body_font

    # ----------------------------------------------------------
    # TABLE
    # ----------------------------------------------------------

    def _add_table(self, table_data: TableData, page_data: PageData):
        if not table_data.rows:
            return

        num_rows = len(table_data.rows)
        num_cols = max((len(r) for r in table_data.rows), default=0)
        if num_rows == 0 or num_cols == 0:
            return

        tbl = self.doc.add_table(rows=num_rows, cols=num_cols)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style = "Table Grid"

        # Calculate column widths
        content_width_pts = (
            page_data.width - page_data.left_margin - page_data.right_margin
        )
        if table_data.col_widths and len(table_data.col_widths) == num_cols:
            total_w = sum(table_data.col_widths)
            col_widths_in = [
                (w / total_w) * (content_width_pts / 72.0)
                for w in table_data.col_widths
            ]
        else:
            col_w = content_width_pts / 72.0 / num_cols
            col_widths_in = [col_w] * num_cols

        for row_idx, row in enumerate(table_data.rows):
            tr = tbl.rows[row_idx]
            for col_idx in range(num_cols):
                cell = tr.cells[col_idx]
                cell_text = row[col_idx] if col_idx < len(row) else ""

                # Set cell width
                try:
                    cell.width = Inches(col_widths_in[col_idx])
                except Exception:
                    pass

                # Clear default paragraph
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(str(cell_text) if cell_text else "")
                run.font.name = self._body_font
                run.font.size = Pt(10)

                # Header row: bold + background
                if row_idx == 0 and table_data.has_header:
                    run.bold = True
                    try:
                        set_cell_shading(cell, table_data.header_bg)
                    except Exception:
                        pass

                # Cell margins
                try:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcMar = parse_xml(
                        f'<w:tcMar {nsdecls("w")}>'
                        f'<w:top w:w="80" w:type="dxa"/>'
                        f'<w:bottom w:w="80" w:type="dxa"/>'
                        f'<w:left w:w="120" w:type="dxa"/>'
                        f'<w:right w:w="120" w:type="dxa"/>'
                        f"</w:tcMar>"
                    )
                    existing = tcPr.find(qn("w:tcMar"))
                    if existing is not None:
                        tcPr.remove(existing)
                    tcPr.append(tcMar)
                except Exception:
                    pass

    # ----------------------------------------------------------
    # IMAGE
    # ----------------------------------------------------------

    def _add_image(self, image_data: ImageData):
        if not image_data.image_bytes:
            return
        try:
            stream = io.BytesIO(image_data.image_bytes)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(
                stream,
                width=Inches(max(0.1, image_data.width_inches)),
                height=Inches(max(0.1, image_data.height_inches)),
            )
            if image_data.caption:
                cp = self.doc.add_paragraph(image_data.caption)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.runs[0].font.size = Pt(9)
                cp.runs[0].italic = True
        except Exception as e:
            print(f"[IMAGE INSERT] {e}")


# =========================================================
# OCR ENGINE (RapidOCR — pure Python, no Tesseract/Poppler)
# =========================================================

def extract_ocr_advanced(pdf_path: str, output_docx: str):
    """OCR using PyMuPDF page rendering + RapidOCR (pure Python)."""
    import numpy as np

    fitz_doc = fitz.open(pdf_path)
    doc = Document()
    _setup_doc_styles(doc)

    for page_idx in range(len(fitz_doc)):
        if page_idx > 0:
            doc.add_page_break()

        page = fitz_doc[page_idx]
        mat = fitz.Matrix(2.0, 2.0)  # 144 DPI
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, 3
        )

        page_text_blocks = []

        if _ocr_engine is not None:
            try:
                result, _ = _ocr_engine(img_array)
                if result:
                    for item in result:
                        # item: [[x1,y1],[x2,y2],[x3,y3],[x4,y4]], text, conf
                        box, text, conf = item[0], item[1], item[2]
                        if conf < 0.3 or not text.strip():
                            continue
                        y_center = (box[0][1] + box[2][1]) / 2
                        page_text_blocks.append((y_center, text.strip()))
            except Exception:
                pass

        if not page_text_blocks:
            # Fallback: use PyMuPDF native text extraction
            text = page.get_text("text").strip()
            for line in text.split("\n"):
                if line.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.font.size = Pt(11)
                    run.font.name = "Arial"
        else:
            # Sort top-to-bottom and write paragraphs
            page_text_blocks.sort(key=lambda x: x[0])
            for _, text in page_text_blocks:
                p = doc.add_paragraph()
                run = p.add_run(text + " ")
                run.font.size = Pt(11)
                run.font.name = "Arial"

    fitz_doc.close()
    doc.save(output_docx)


def _setup_doc_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)




# =========================================================
# HYBRID EXTRACTION
# =========================================================

def extract_hybrid(pdf_path: str, output_docx: str):
    """Auto-detect best extraction method per document."""
    doc = fitz.open(pdf_path)
    total_text = ""
    total_pages = len(doc)
    selectable_pages = 0

    for page in doc:
        pt = page.get_text().strip()
        total_text += pt
        if len(pt) > 20:
            selectable_pages += 1

    doc.close()

    mostly_scanned = selectable_pages < total_pages * 0.5

    if len(total_text.strip()) < 50 or mostly_scanned:
        print("[HYBRID] Falling back to OCR")
        extract_ocr_advanced(pdf_path, output_docx)
    else:
        print("[HYBRID] Using native layout extraction")
        extractor = AdvancedPDFExtractor(pdf_path)
        pages_data = extractor.extract_all()
        extractor.close()
        builder = AdvancedDocxBuilder(pages_data)
        builder.build(output_docx)


# =========================================================
# FLASK ROUTES
# =========================================================

@app.route("/health")
def health():
    return jsonify({
        "status": "online",
        "server": "HULK BUSTER ULTRA",
        "version": "3.0",
        "features": [
            "native", "ocr", "layout", "hybrid",
            "tables", "images", "hyperlinks",
            "headers_footers", "lists", "headings",
            "multi_column", "rtl", "font_mapping",
            "relative_heading_detection",
            "superscript_subscript",
            "code_blocks", "alignment_detection",
            "margin_estimation",
        ],
    })


@app.route("/convert", methods=["POST", "OPTIONS"])
def convert():
    if request.method == "OPTIONS":
        return ("", 204)

    workspace = None
    try:
        file = request.files.get("file")
        validate_pdf(file)
        mode = request.form.get("convert_mode", "hybrid")
        workspace = tempfile.mkdtemp(prefix="hulk_ultra_")
        filename = secure_filename(file.filename)
        pdf_path = os.path.join(workspace, filename)
        output_docx = os.path.join(workspace, f"{Path(filename).stem}.docx")
        file.save(pdf_path)

        if mode == "native":
            converter = Converter(pdf_path)
            try:
                converter.convert(output_docx, start=0, end=None)
            finally:
                converter.close()

        elif mode == "ocr":
            extract_ocr_advanced(pdf_path, output_docx)

        elif mode == "layout":
            extractor = AdvancedPDFExtractor(pdf_path)
            pages_data = extractor.extract_all()
            extractor.close()
            builder = AdvancedDocxBuilder(pages_data)
            builder.build(output_docx)

        elif mode == "hybrid":
            extract_hybrid(pdf_path, output_docx)

        else:
            raise Exception(f"Invalid convert_mode: {mode}")

        return send_file(
            output_docx,
            as_attachment=True,
            download_name=f"{Path(filename).stem}.docx",
            mimetype=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
        )

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        if workspace:
            safe_cleanup(workspace)
        gc.collect()


@app.route("/analyze", methods=["POST", "OPTIONS"])
def analyze():
    if request.method == "OPTIONS":
        return ("", 204)

    workspace = None
    try:
        file = request.files.get("file")
        validate_pdf(file)
        workspace = tempfile.mkdtemp(prefix="hulk_analyze_")
        filename = secure_filename(file.filename)
        pdf_path = os.path.join(workspace, filename)
        file.save(pdf_path)

        doc = fitz.open(pdf_path)
        metadata = {
            "filename": filename,
            "pages": len(doc),
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "subject": doc.metadata.get("subject", ""),
            "creator": doc.metadata.get("creator", ""),
            "producer": doc.metadata.get("producer", ""),
            "creation_date": doc.metadata.get("creationDate", ""),
            "modification_date": doc.metadata.get("modDate", ""),
            "encrypted": doc.is_encrypted,
            "file_size": os.path.getsize(pdf_path),
        }

        analyzer = FontSizeAnalyzer()
        pages_info = []
        total_text_len = 0
        total_images = 0
        total_tables = 0

        for pn in range(len(doc)):
            page = doc[pn]
            analyzer.feed_page(page)
            ti = {
                "page": pn + 1,
                "width": page.rect.width,
                "height": page.rect.height,
                "rotation": page.rotation,
                "text_length": len(page.get_text()),
                "image_count": len(page.get_images()),
                "table_count": 0,
            }
            try:
                ti["table_count"] = len(page.find_tables())
                total_tables += ti["table_count"]
            except Exception:
                pass
            pages_info.append(ti)
            total_text_len += ti["text_length"]
            total_images += ti["image_count"]

        doc.close()

        body_size = analyzer.body_size()
        if total_text_len < 100:
            recommended_mode = "ocr"
        elif total_images > 5 and total_text_len < 500:
            recommended_mode = "hybrid"
        else:
            recommended_mode = "layout"

        return jsonify({
            "metadata": metadata,
            "pages": pages_info,
            "statistics": {
                "total_text_length": total_text_len,
                "total_images": total_images,
                "total_tables": total_tables,
                "detected_body_font_size": body_size,
                "recommended_mode": recommended_mode,
                "is_scanned": total_text_len < 100,
            },
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

    finally:
        if workspace:
            safe_cleanup(workspace)
        gc.collect()


# =========================================================
# START
# =========================================================

if __name__ == "__main__":
    print("=" * 55)
    print("  HULK BUSTER ULTRA SERVER v3.0")
    print("=" * 55)
    print("  POST /convert   — Convert PDF → Word (.docx)")
    print("  POST /analyze   — Analyze PDF metadata")
    print("  GET  /health    — Server health check")
    print("-" * 55)
    print("  Modes: native | ocr | layout | hybrid (default)")
    print("  http://localhost:5000")
    print("=" * 55)

    serve(
        app,
        host="0.0.0.0",
        port=5000,
        threads=16,
        connection_limit=1000,
        channel_timeout=180,
    )